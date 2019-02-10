import os
import pymysql, collections
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Nodes
from .forms import UploadFileForm, UserForm
# -*- coding: utf-8 -*-

from _ast import For

def auth_login(request):
    return render(request, 'website/login.html')

def registration(request):
    if (request.user.is_authenticated()):
        return render(request, 'website/registration.html')
    else:
	return render(request, 'website/login.html')

def register(request):
    form_class = UserForm
    if(request.method == "POST"):
        form = form_class(request.POST)
        if form.is_valid():

            user = form.save(commit=False)

            #cleaned normalized data
            username = form.cleaned_data['username']
            email = form.cleaned_data['email']
            password = form.cleaned_data['password']
            user.set_password(password)
            user.save()
            return redirect('/CiberC/')
    else:
        return render(request, 'website/login.html', {'error_message': "Registration Failed"})

def main(request):
    if (request.user.is_authenticated()):
        return render(request, 'website/main.html')
    else:
	return render(request, 'website/login.html')

def login_validation(request):
    if(request.method == "POST"):
        #form = UserForm(request.POST)
        #if(form.is_valid()):
        username = request.POST['username']
        password = request.POST['password']
        email = request.POST['username']

        #return User object if credentials are correct
        user = authenticate(username=username, password=password)

        if user is not None:
            if user.is_active:
                login(request, user)
                return redirect('/CiberC/main/')
    return render(request, 'website/login.html', {'error_message': "Username or Password incorrect"})


def auth_logout(request):
    logout(request)
    return redirect('/CiberC/')


def upload(request):
    if (request.user.is_authenticated()):
	return render(request,'website/upload.html')
    else:
	return render(request, 'website/login.html')

# handle_uploaded_file() -> Not a view
def handle_uploaded_file(file, filename):
    if not os.path.exists('file/'):
        os.mkdir('file/')

    with open('file/'+ filename, 'wb+') as destination:

        for chunk in file.chunks():
            destination.write(chunk)

def upload_file(request):
    if (request.user.is_authenticated()):
        if request.method == 'POST':
            handle_uploaded_file(request.FILES['file'], str(request.FILES['file']))
            acum=load_date('/var/www/html/CiberC/file/'+str(request.FILES['file']))
            if acum==0:
	        return render(request, 'website/upload.html', {'error_message': "Error at uploading the File"})
            else:
            	return render(request, 'website/upload.html', {'success_message': "File uploaded"})
        return render(request, 'website/upload.html', {'error_message': "Error at uploading the File"})
    else:
	return render(request, 'website/login.html')


def download(request):
    if (request.user.is_authenticated()):
        all_nodes = Nodes.objects.all()

        if os.path.exists('/var/www/html/CiberC/file/isistopology.txt'):
            return render(request, 'website/download.html', {'success_message': 'Please insert the name or ip address of the node:'})

        if all_nodes.exists():
            return render(request, 'website/download.html', {'success_message': 'Please insert the name or ip address of the node:'})
        return render(request, 'website/download.html', {'error_message': "No Data found. Please upload the isistopology.txt in the Upload section."})
    else:
	return render(request, 'website/login.html')


def download_file(request):
    if (request.user.is_authenticated()):
        if (request.method == "POST"):
            nodo = request.POST['node']
            texto=conexas(nodo)
    	    document = Document()
	    number=0
 	    d=document.add_heading('CiberC-Nodos', 0)
	    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
	    p=document.add_paragraph('The selected node is: ')
            p.add_run(nodo).bold = True
	    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	    if len(texto)>1:
            	p = document.add_paragraph('The following groups of affected nodes were created: ')
            	p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
            	p = document.add_paragraph('No affected nodes')
           	p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for components in texto:
            	if number>0:
                	s=document.add_heading("Group #%i:" % (number), level=1)
                	s.alignment = WD_ALIGN_PARAGRAPH.CENTER
                	p = document.add_paragraph("%s" % (list(components).__str__().replace('[','').replace(']','')))
                	p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            	number += 1
    	    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    	    response['Content-Disposition'] = 'attachment; filename=download.docx'
    	    document.save(response)
    	    return response

        return render(request, 'website/download.html', {'success_message': 'Please insert the name or ip address of the node:'})
    else:
	return render(request, 'website/login.html')


def view(request):
    if (request.user.is_authenticated()):
        all_nodes = Nodes.objects.all()

        if os.path.exists('/var/www/html/CiberC/file/isistopology.txt'):
            return render(request, 'website/view.html', {'success_message': 'Please insert the name or ip addres of the node:'})
        print(all_nodes)
        if all_nodes.exists():
            return render(request, 'website/view.html',
                     		  {'success_message': 'Please insert the name or ip address of the node:'})
        return render(request, 'website/view.html', {'error_message': "No Data found. Please upload the isistopology.txt in the Upload section."})
    else:
	return render(request, 'website/login.html')


def view_graph(request):
    if (request.user.is_authenticated()):
        if (request.method == "POST"):
            nodo = request.POST['node']
            level = request.POST['level']
            red = request.POST['red']
            link = request.POST['link']
            conn = pymysql.connect(
                host="localhost", port=3306, user="root",
                passwd="hotmail003", db="red"
            )
            print(level)
            print(link)
            texto = ""
            textono = ""
            nodo=convertir(nodo)

            try:
                list = []
		listIp= []
                with conn.cursor() as cursor:
                    sql0 = "SELECT * FROM `isis`"
                    cursor.execute(sql0)
                    for row in cursor:
                        ip = row[1].rstrip("\r")
                        listIp.append(ip)
			nodo1 = row[0].rstrip("\r")
                        list.append(nodo1)

                with conn.cursor() as cursor:
                    sql1 = "SELECT * FROM `isis` INNER JOIN provider USING(Hostname)"
                    sql2 = "SELECT * FROM `isis` INNER JOIN provideredge USING(Hostname)"
                    sql3 = "SELECT * FROM `isis` INNER JOIN border USING(Hostname)"
                    sql = sql1 + " UNION " + sql2 + " UNION " + sql3 + " ORDER BY Hostname ASC"
                    cursor.execute(sql)
                    result=cursor.fetchall()
                    listaT = []
                    listaT.append(nodo)
                    lista = []
                    lista1 = []
                    lista2 = []
                    edge= []
                    edge1= []
                    for row in result:
                        nodo1=row[0].rstrip("\r")
                        lista.append(row)
                        palabras = row[2].split(".")
                        if level == "4":
                            palabras1 = palabras[0].split("-")
                            if (len(palabras1) >= 2):
                                if (palabras1[1] == "COR"):
                                    palabras[0] = palabras[0] + "E"
                            if not palabras[0] in listaT:
                                lista1.append(palabras[0])
                                listaT.append(palabras[0])
                            if link == "1":
                                edgeT = "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(list.index(palabras[0])) + ", \"arrows\": \"to\" },\n"
                                edgeT1 = "{\"from\": " + str(list.index(palabras[0]))  + ", \"to\": " + str(list.index(nodo1))  + ", \"arrows\": \"to\" },\n"
                                if not edgeT in edge:
                                    if not edgeT1 in edge:
                                        edge.append(edgeT)
                                        textono = textono + edgeT
                            else:
                                textono = textono + "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                list.index(palabras[0])) + ", \"label\": \"" + str(row[4]) + "\", \"arrows\": \"to\" },\n"

                        else:
                            if nodo1==nodo:
                                palabras1 = palabras[0].split("-")
                                if (len(palabras1) >= 2):
                                    if (palabras1[1] == "COR"):
                                        palabras[0] = palabras[0] + "E"

                                if level == "1" or level == "3" or level == "2":
                                    if not palabras[0] in listaT:
                                        lista1.append(palabras[0])
                                        listaT.append(palabras[0])
                                if link == "1":
                                    edgeT = "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) +", \"arrows\": \"to\" },\n"
                                    edgeT1 = "{\"from\": " + str(list.index(palabras[0])) + ", \"to\": " + str(
                                        list.index(nodo1)) +", \"arrows\": \"to\" },\n"
                                    if not edgeT in edge:
                                        if not edgeT1 in edge:
                                            edge.append(edgeT)
                                            textono = textono + edgeT
                                else:
                                    textono = textono + "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) +", \"label\": \""+str(row[4])+"\", \"arrows\": \"to\" },\n"
                    for k in lista:
                        nodo1 = k[0].rstrip("\r")
                        if nodo1 in lista1:
                            palabras = k[2].split(".")
                            palabras1 = palabras[0].split("-")
                            if (len(palabras1) >= 2):
                                if (palabras1[1] == "COR"):
                                    palabras[0] = palabras[0] + "E"
                            if level == "3":
                                if not palabras[0] in listaT:
                                    lista2.append(palabras[0])
                            if level=="4":
                                if not palabras[0] in listaT:
                                    listaT.append(palabras[0])
                                if link == "1":
                                    edgeT = "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) +", \"arrows\": \"to\" },\n"
                                    edgeT1 = "{\"from\": " + str(list.index(palabras[0])) + ", \"to\": " + str(
                                        list.index(nodo1)) + ", \"arrows\": \"to\" },\n"
                                    if not edgeT in edge:
                                        if not edgeT1 in edge:
                                            edge.append(edgeT)
                                            textono = textono + edgeT
                                else:
                                    textono = textono + "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) + ", \"label\": \"" + str(k[4]) +"\", \"arrows\": \"to\" },\n"

                            if level == "3" or level == "2" :
                                if not palabras[0] in listaT:
                                    listaT.append(palabras[0])
                                if link == "1":
                                    edgeT = "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) + ", \"arrows\": \"to\" },\n"
                                    edgeT1 = "{\"from\": " + str(list.index(palabras[0])) + ", \"to\": " + str(
                                        list.index(nodo1)) + ", \"arrows\": \"to\" },\n"
                                    if not edgeT in edge:
                                        if not edgeT1 in edge:
                                            edge.append(edgeT)
                                            textono = textono + edgeT
                                else:
                                    textono = textono + "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) + ", \"label\": \"" + str(k[4]) +"\", \"arrows\": \"to\" },\n"

                    if level == "3":
                        for j in lista:
                            nodo1 = j[0].rstrip("\r")
                            if nodo1 in lista2:
                                palabras = j[2].split(".")
                                palabras1 = palabras[0].split("-")
                                if (len(palabras1) >= 2):
                                    if (palabras1[1] == "COR"):
                                        palabras[0] = palabras[0] + "E"
                                if level == "3":
                                    if not palabras[0] in listaT:
                                        listaT.append(palabras[0])
                                if link == "1":
                                    edgeT = "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) + ", \"arrows\": \"to\" },\n"
                                    edgeT1 = "{\"from\": " + str(list.index(palabras[0])) + ", \"to\": " + str(
                                        list.index(nodo1)) + ", \"arrows\": \"to\" },\n"
                                    if not edgeT in edge:
                                        if not edgeT1 in edge:
                                            edge.append(edgeT)
                                            textono = textono + edgeT
                                else:
                                    textono = textono + "{\"from\": " + str(list.index(nodo1)) + ", \"to\": " + str(
                                        list.index(palabras[0])) + ", \"label\": \"" + str(j[4]) + "\", \"arrows\": \"to\" },\n"

                    for t in listaT:
                        if t[-3] == "P":
                            Tipo = "P"
                        if t[-3] == "E":
                            Tipo = "E"
                        if t[-3] == "R":
                            Tipo = "R"
                        if t[-3] == "B":
                            Tipo = "B"
                        if t == nodo:
                            Tipo = "C"
                        if t in list:
                            texto = texto + "{id:" + str(list.index(t)) + ", \"label\": \"" + t + "\",\"title\": \"" + str(listIp[list.index(t)]) + "\",\"group\": \"" + Tipo + "\"},\n"
            finally:
                # Close connection.
                conn.close()
            if red=="1":
                return render(request, 'website/GrafoR.html', {'nodos': texto, 'edge': textono})
            else:
                return render(request, 'website/Grafo.html', {'nodos': texto,'edge': textono})
    else:
	return render(request, 'website/login.html')


def load_date(documento):
    conn = pymysql.connect(
        host="localhost", port=3306, user="root",
        passwd="hotmail003", db="red"
    )
    fileShow = open(documento, "r")
    lineas = fileShow.readlines()
    acum3 = 0;
    hostname = []
    ip = []
    ipGlobal = []
    metric = []
    extended = []
    vecinosGlobal = []
    metricGlobal = []
    extendedGlobal = []
    vecinos = []
    for i in lineas:
        palabras = i.split(" ")
        if palabras[0] == "Connection":
            vecinosGlobal.append(vecinos)
            metricGlobal.append(metric)
            extendedGlobal.append(extended)
            if (ip):
                ipGlobal.append(ip)
            else:
                ip.append("0")
                ipGlobal.append(ip)
            ip = []
            vecinos = []
            metric = []
            extended = []
        if len(palabras) >= 3:
            if palabras[2] == "Hostname:":
                # valid if the node have neighbor
                if acum3 == 1:
                    vecinos.append("0")
                    acum3 = 0
                palabras1 = palabras[3].split("\n")
                hostname.append(palabras1[0])
                #valid if the list hostname have one node
                if len(hostname) == 1:
                    vecinos = []
                    metric = []
                    extended = []
                    ip = []
                else:
                    vecinosGlobal.append(vecinos)
                    metricGlobal.append(metric)
                    extendedGlobal.append(extended)
                    # valid if the node have ip
                    if (ip):
                        ipGlobal.append(ip)
                    else:
                        ip.append("0")
                        ipGlobal.append(ip)
                    ip = []
                    vecinos = []
                    metric = []
                    extended = []
        if len(palabras) >= 5:
            # valid if the line have the word "IP Address"
            if palabras[2] == "IP" and palabras[3] == "Address:":
                palabras1 = palabras[-1].split("\n")
                ip.append(palabras1[0])
            # valid if the line have the word "Neighbor"
            if palabras[4] == "Neighbor":
                palabras1 = palabras[-1].split("\n")
                vecinos.append(palabras1[0])
                acum3 = 0
            # valid if the node have neighbor
            if acum3 == 1:
                vecinos.append("0")
                acum3 = 0
            # valid if the line have the word "Metric"
            if palabras[2] == "Metric:":
                acum3 = 1
                metric.append(palabras[3])
                palabras1 = palabras[-1].split("\n")
                extended.append(palabras1[0])

    acum = 0;
    #close the file
    fileShow.close()
    #valid if the file is correct
    if len(hostname)<1:
        return 0
    #if the file is correct then we delete all the information from the database
    else:
        try:
            sql = "DELETE FROM `border` "
            sql1 = "DELETE FROM `provider` "
            sql2 = "DELETE FROM `provideredge` "
            sql3 = "DELETE FROM `isis`"
            with conn.cursor() as cursor:
                cursor.execute(sql)
                conn.commit()
                cursor.execute(sql1)
                conn.commit()
                cursor.execute(sql2)
                conn.commit()
            with conn.cursor() as cursor:
                cursor.execute(sql3)
                conn.commit()

            for i in hostname:
                # insert the isis nodes in the database
                sql = "INSERT INTO `isis`( `Hostname`, `Ip`) VALUES ('" + hostname[acum] + "','" + ipGlobal[acum][0] + "') "
                for m in ipGlobal[acum]:
                    if not m==ipGlobal[acum][0]:
                        sql1 = "INSERT INTO `Ip`( `Hostname`, `Ip`) VALUES ('" + hostname[acum] + "','" + m + "') "
                        with conn.cursor() as cursor:
                            cursor.execute(sql1)
                            conn.commit()
                print(acum)
                with conn.cursor() as cursor:
                     cursor.execute(sql)
                     conn.commit()
                acum1 = 0;
                # insert the neighbors nodes in the database
                with conn.cursor() as cursor:
                    for j in extendedGlobal[acum]:
                        palabras1 = extendedGlobal[acum][acum1].split(".")
                        cadena = list(palabras1[0])

                        if not cadena[1] == "-":
                            palabras2 = palabras1[0].split("-")
                            cadena1 = list(palabras2[0])
                        else:
                            palabras2 = palabras1[0]
                            cadena1 = list(palabras2)

                        if cadena1[-3] == "P":
                            sql = "INSERT INTO `provider`( `Hostname`, `Neighbor`, `Ip`, `Metrica`) VALUES ('" + hostname[
                                acum] + "','" + extendedGlobal[acum][acum1] + "','" + vecinosGlobal[acum][acum1] + "','" + \
                                  metricGlobal[acum][acum1] + "') "
                        if cadena1[-3] == "E":
                            sql = "INSERT INTO `provideredge`( `Hostname`, `Neighbor`, `Ip`, `Metrica`) VALUES ('" + \
                                  hostname[acum] + "','" + extendedGlobal[acum][acum1] + "','" + vecinosGlobal[acum][
                                      acum1] + "','" + metricGlobal[acum][acum1] + "') "
                        if cadena1[-3] == "B":
                            sql = "INSERT INTO `border`( `Hostname`, `Neighbor`, `Ip`, `Metrica`) VALUES ('" + hostname[
                                acum] + "','" + extendedGlobal[acum][acum1] + "','" + vecinosGlobal[acum][acum1] + "','" + \
                                  metricGlobal[acum][acum1] + "') "
                        if cadena1[-3] == "R":
                            sql = "INSERT INTO `border`( `Hostname`, `Neighbor`, `Ip`, `Metrica`) VALUES ('" + hostname[
                                acum] + "','" + extendedGlobal[acum][acum1] + "','" + vecinosGlobal[acum][acum1] + "','" + \
                                  metricGlobal[acum][acum1] + "')"
                        cursor.execute(sql)
                        acum1 = acum1 + 1
                        conn.commit()
                acum = acum + 1
        finally:
            # Close connection.
            conn.close()

def convertir(nodo):
        conn = pymysql.connect(
            host="localhost", port=3306, user="root",
            passwd="hotmail003", db="red"
        )
        nod = nodo.split(".")
        bandera=0
        if (nod >= 2):
            bandera=1
        try:
            with conn.cursor() as cursor:
                sql0 = "SELECT * FROM `isis`"
                cursor.execute(sql0)
                for row in cursor:
                    nodo1 = row[0].rstrip("\r")
                    nodo2 = row[1].rstrip("\r")
                    if nodo2 == nodo:
                        bandera = 0
                        nodo = nodo1

            if bandera==1:
                with conn.cursor() as cursor:
                    sql0 = "SELECT * FROM `Ip`"
                    cursor.execute(sql0)
                    for row in cursor:
                        nodo2 = row[1].rstrip("\r")
                        if nodo2 == nodo:
                            nodo1 = row[0].rstrip("\r")
                            bandera = 0
                            nodo = nodo1
            if bandera==1:
                with conn.cursor() as cursor:
                    sql1 = "SELECT * FROM `isis` INNER JOIN provider USING(Hostname)"
                    sql2 = "SELECT * FROM `isis` INNER JOIN provideredge USING(Hostname)"
                    sql3 = "SELECT * FROM `isis` INNER JOIN border USING(Hostname)"
                    sql = sql1 + " UNION " + sql2 + " UNION " + sql3 + " ORDER BY Hostname ASC"
                    cursor.execute(sql)
                    for row in cursor:
                        nodo2 = row[3].rstrip("\r")
                        if nodo2 == nodo:
                            nodo1 = row[2].rstrip("\r")
                            nod2 = nodo1.split(".")
                            bandera = 0
                            nodo = nod2[0]
        finally:
            # Close connection.
            conn.close()
        return nodo

def connected_components(nodes, diccionario):
    # List of connected components found. The order is random.
    result = []

    # Make a copy of the set, so we can modify it.
    nodes = set(nodes)

    # Iterate while we still have nodes to process.
    while nodes:

        # Get a random node and remove it from the global set.
        n = nodes.pop()

        # This set will contain the next group of nodes connected to each other.
        group = {n}

        # Build a queue with this node in it.
        queue = [n]

        # Iterate the queue.
        # When it's empty, we finished visiting a group of connected nodes.
        while queue:
            # Consume the next item from the queue.
            n = queue.pop(0)
            # Fetch the neighbors.

            neighbors = diccionario.get(n)

            if neighbors==None:
                break
            if len(neighbors)>0:
                # Remove the neighbors we already visited.
                neighbors.difference_update(group)

                # Remove the remaining nodes from the global set.
                nodes.difference_update(neighbors)

                # Add them to the group of connected nodes.
                group.update(neighbors)
                #print(group)
                # Add them to the queue, so we visit them in the next iterations.
                queue.extend(neighbors)

        # Add the group to the list of groups.
        result.append(group)
    # Return the list of groups.

    return result

def conexas(nodo):
    conn = pymysql.connect(
        host="localhost", port=3306, user="root",
        passwd="hotmail003", db="red"
    )
    nodo=convertir(nodo)
    print(nodo)
    try:
        nodes = set()
        list = []
        with conn.cursor() as cursor:
            sql0 = "SELECT * FROM `isis`"
            cursor.execute(sql0)
            for row in cursor:
                nodo1 = row[0].rstrip("\r")
                list.append(nodo1)

        with conn.cursor() as cursor:
            sql1 = "SELECT * FROM `isis` INNER JOIN provider USING(Hostname)"
            sql2 = "SELECT * FROM `isis` INNER JOIN provideredge USING(Hostname)"
            sql3 = "SELECT * FROM `isis` INNER JOIN border USING(Hostname)"
            sql = sql1 + " UNION " + sql2 + " UNION " + sql3 + " ORDER BY Hostname ASC"
            cursor.execute(sql)
            result = cursor.fetchall()
            lista = []
            lista1 = []
            diccionario = dict()
            listaT = []
            listaT.append(nodo)
            for row in result:
                nodo1 = row[0].rstrip("\r")
                lista.append(row)
                palabras = row[2].split(".")
                palabras1 = palabras[0].split("-")
                if (len(palabras1) >= 2):
                    if (palabras1[1] == "COR"):
                        palabras[0] = palabras[0] + "E"
                        lista1.append(palabras[0])
                if not palabras[0] in listaT:
                    listaT.append(palabras[0])
                if str(nodo1) != nodo and str(palabras[0]) != nodo:
                    listaVal = set()
                    if str(nodo1) in diccionario:

                        for x in diccionario.get(str(nodo1)):
                            if not x in listaVal:
                                listaVal.add(x)
                        listaVal.add(str(palabras[0]))
                        diccionario[str(nodo1)] = listaVal
                    else:
                        listaVal.add(str(palabras[0]))
                        diccionario[str(nodo1)] = listaVal
                for k in lista:
                    nodo1 = k[0].rstrip("\r")
                    listaVal = set()
                    if nodo1 in lista1:
                        palabras = k[2].split(".")
                        palabras1 = palabras[0].split("-")
                        if (len(palabras1) >= 2):
                            if (palabras1[1] == "COR"):
                                palabras[0] = palabras[0] + "E"
                        if not palabras[0] in listaT:
                            listaT.append(palabras[0])
                        if str(nodo1) != nodo and str(palabras[0]) != nodo:
                            if str(nodo1) in diccionario:
                                for x in diccionario.get(str(nodo1)):
                                    if not x in listaVal:
                                        listaVal.add(x)
                                listaVal.add(str(palabras[0]))
                                diccionario[str(nodo1)] = listaVal
                            else:
                                listaVal.add(str(palabras[0]))
                                diccionario[str(nodo1)] = listaVal
            for t in listaT:
                if t in list:
                    if not t == nodo:
                        nodes.add(t)

	    return  connected_components(nodes, diccionario)
    finally:
        # Close connection.
        conn.close()
def modules(request):
    if (request.user.is_authenticated()):
        return render(request, 'website/modules.html')
    else:
	return render(request, 'website/login.html')

def convertModules(request):
    if request.method == 'POST':
        if request.POST.get('commands') == '':
            return render(request, 'website/modules.html', {'error_message': "Please insert the commands lines to proceed the conversion."})

        command_text = request.POST.get('commands')
        option = request.POST.get('opt')
        #commands Vfis
        document = Document()
        number = 0
        # Title of the document
        d = document.add_heading('CiberC-Modules', 0)
        # center the title
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Create one paragraph
        p = document.add_paragraph('The selected Module is: ')
        command_text_iosXR=""
        if request.POST.get('option') == 'Vfis':
            # change the name of the node to bold and center the text
            p.add_run('Vfis').bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # condition to validate if there are groups of affected nodes
            command_text_iosXR = vfis(command_text)
            if command_text_iosXR == "The commands entered are incorrect ":
                return render(request, 'website/modules.html', {'ios': command_text,
                                                         'iosXR': command_text_iosXR})
            else:
                if option == '1':
                    p = document.add_paragraph(command_text_iosXR)
                    # 'for' the groups of affected nodes and saved them in the document
                    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = 'attachment; filename=Vfis.docx'
                    document.save(response)
                    return response
                elif option == '2':
                    return render(request, 'website/modules.html', {'ios': command_text,
                                                                    'iosXR': command_text_iosXR})
        # commands Xconnects
        elif request.POST.get('option') == 'Xconnects':
            # change the name of the node to bold and center the text
            p.add_run('Xconnects').bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # condition to validate if there are groups of affected nodes
            command_text_iosXR = xconnects(command_text)
            if command_text_iosXR == "The commands entered are incorrect ":
                return render(request, 'website/modules.html', {'ios': command_text,
                                                         'iosXR': command_text_iosXR})
            else:
                if option == '1':
                    p = document.add_paragraph(command_text_iosXR)
                    # 'for' the groups of affected nodes and saved them in the document
                    response = HttpResponse(
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = 'attachment; filename=Xconnects.docx'
                    document.save(response)
                    return response
                elif option == '2':
                    return render(request, 'website/modules.html', {'ios': command_text,
                                                                    'iosXR': command_text_iosXR})
        # commands Prefix-list
        elif request.POST.get('option') == 'Prefix-List':
            # change the name of the node to bold and center the text
            p.add_run('Prefix-List').bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # condition to validate if there are groups of affected nodes
            command_text_iosXR = prefix(command_text)
            if command_text_iosXR == "The commands entered are incorrect ":
                return render(request, 'website/modules.html', {'ios': command_text,
                                                         'iosXR': command_text_iosXR})
            else:
                if option == '1':
                    p = document.add_paragraph(command_text_iosXR)
                    # 'for' the groups of affected nodes and saved them in the document
                    response = HttpResponse(
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = 'attachment; filename=Prefix-LIst.docx'
                    document.save(response)
                    return response
                elif option == '2':
                    return render(request, 'website/modules.html', {'ios': command_text,
                                                                    'iosXR': command_text_iosXR})

        return render(request, 'website/modules.html', {'ios': command_text,
                                                        'iosXR': command_text_iosXR})
def xconnects(commands_ios):
    commands_iosXR=""
    palabra = commands_ios.split("\n")
    for i in palabra:
        palabras = i.split(" ")
        palabras = ' '.join(palabras).split()
        if len(palabras)>2:
            if palabras[1]=="ac" and palabras[3]=="VLAN)":
                palabras1 = palabras[2].split(":")
                vlan = palabras1[0].strip('Vl')
                text = "l2vpn bridge group CNT bridge-domain " + vlan + "\n"
                ' '.join(palabras).split()
                ip = palabras[-2].split(":")
                text = text + "neighbor " + ip[0] + "pw-id " + ip[1] + "\n" + "root" + "\n" + "!"+ "\n"
                commands_iosXR = commands_iosXR + text
            else:
                return "The commands entered are incorrect "
    return commands_iosXR

def vfis(commands_ios):
    commands_iosXR = ""
    palabra = commands_ios.split("\n")
    text=""
    bandera=1
    for i in palabra:
        palabras = i.split(" ")
        palabras = ' '.join(palabras).split()
        if len(palabras)>2:
            text = ""
            if palabras[0] == "l2":
                if bandera==0:
                    text = "root" + "\n" + "!"+ "\n"
                    bandera=1
                domain = palabras[2]
            if palabras[0] == "vpn":
                id= palabras[2]
                text = "l2vpn bridge group CNT bridge-domain " + palabras[2] + "\n"
                text = text + "vfi " + domain+ "\n"
            if palabras[0] == "neighbor":
                text = "neighbor "+palabras[1]+ " pw-id "+ id+ "\n"
                bandera=0
        commands_iosXR = commands_iosXR + text
    if commands_iosXR=="":
        return "The commands entered are incorrect "
    else:
        commands_iosXR = commands_iosXR + "root" + "\n" + "!"+ "\n"
    return commands_iosXR

def prefix(commands_ios):
    commands_iosXR = ""
    diccionario = collections.OrderedDict()
    palabra = commands_ios.split("\n")
    for i in palabra:
        listaVal = []
        palabras = i.split(" ")
        palabras = ' '.join(palabras).split()
        if len(palabras) > 2:
            print(palabras)
            if palabras[1] == "prefix-list" and palabras[-2] == "permit":
                if str(palabras[2]) in diccionario:
                    for x in diccionario.get(str(palabras[2])):
                        if not x in listaVal:
                            listaVal.append(x)
                    listaVal.append(str(palabras[-1]))
                    diccionario[str(palabras[2])] = listaVal
                else:
                    listaVal.append(str(palabras[-1]))
                    diccionario[str(palabras[2])] = listaVal
            else:
                return "The commands entered are incorrect "
    for m in diccionario:
        commands_iosXR=commands_iosXR + "prefix-set pfx_"+m+"_p1_permit"+ "\n"
        for x in diccionario[m]:
            commands_iosXR = commands_iosXR + x + "\n"
        commands_iosXR = commands_iosXR + "end-set"  + "\n"
    return commands_iosXR

